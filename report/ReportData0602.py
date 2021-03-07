# -*- coding: utf-8 -*-
import pandas as pd
import requests
from bs4 import BeautifulSoup
import re
import time
import datetime as dt
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def getDict():        
    url_list = ['https://cn.investing.com/indices/major-indices',
                'https://cn.tradingview.com/markets/indices/quotes-snp/',
                'https://cn.investing.com/rates-bonds/world-government-bonds',
                'https://cn.investing.com/currencies/us-dollar-index',
                'https://cn.investing.com/indices/china-a50'              
                ]
    
    url_list_his = ['https://cn.investing.com/indices/us-30-historical-data',
                   'https://cn.investing.com/indices/us-spx-500-historical-data',
                   'https://cn.investing.com/indices/nasdaq-composite-historical-data',
                   'https://cn.investing.com/indices/uk-100-historical-data',
                   'https://cn.investing.com/indices/france-40-historical-data',
                   'https://cn.investing.com/indices/germany-30-historical-data',
                   'https://cn.investing.com/indices/spain-35-historical-data',
                   'https://cn.investing.com/indices/it-mib-40-historical-data',
                   'https://cn.investing.com/indices/volatility-s-p-500-historical-data',
                   'https://cn.investing.com/rates-bonds/u.s.-10-year-bond-yield-historical-data',
                   'https://cn.investing.com/currencies/us-dollar-index-historical-data',
                   'https://cn.investing.com/indices/china-a50-historical-data'
                   ]
    
    kv = {'user-agent': 'Mozilla/5.0',
          'Connection':'close'}
    
    index = [u'道指',
             u'标普500指数',
             u'纳指',
             u'英国富时100指数',
             u'法国CAC40指数',
             u'德国DAX指数',
             u'西班牙IBEX35指数',
             u'意大利富时MIB指数',
             u'芝加哥期权交易所市场波动率指数（VIX）'
             ]
    
    ref = ['pid-169',
           'pid-166',
           'pid-14958',
           'pid-27',
           'pid-167',
           'pid-172',
           'pid-174',
           'pid-177',
           'pid-44336'
          ]
    
    mark = [u'；',
            u'；',
            u'。',
            u'；',
            u'；',
            u'；',
            u'；',
            u'。'
           ]  
    
    sector = {'SP:S5COND': u'可选消费',
              'SP:S5CONS': u'必需消费',
              'SP:S5HLTH': u'医疗保健',
              'SP:S5INDU': u'工业',
              'SP:S5INFT': u'信息技术',
              'SP:S5MATR': u'材料',
              'SP:S5REAS': u'房地产',
              'SP:S5TELS': u'通讯服务',
              'SP:S5UTIL': u'公共事业',
              'SP:SPF': u'金融',
              'SP:SPN': u'能源',
              }
    
    TBlist = ['23697',
              '23698',
              '23699',
              '23700',
              '23701',
              '23702',
              '23703',
              '23704',
              '23705',
              '23706']

    
    return url_list, url_list_his, kv, index, ref, mark, sector, TBlist
  

def weekEnd(rtday, url_list_his, getSoup, kv):
    if rtday.weekday() == 4:
        print ('----------报告日为周五，计算周度数据------------')
        prewk_last_list = []
        l = len(url_list_his)
        for i in range(l):
            print ('----------正在爬取第%d个周度数据，共%d个周度数据'%(i+1, l))
            last_weekday = rtday - dt.timedelta(days=7)
            last_date = str(last_weekday.year)+u'年'+str(last_weekday.month)+u'月'+str(last_weekday.day)+u'日'
            url = url_list_his[i]
            soup = getSoup(url, kv)
            t = soup.find_all('table', {'class': 'genTbl closedTbl historicalTbl'})
            for i in t:
                j = i.text
            k = j.split()
            try:
                pre_wk_last = k[k.index(last_date)+1]
            except:
                print('!!!读取错误，录入上个指数数据，注意修改prewk_last_list的第%d个数据!!!地址为%s'%(i, url))
            prewk_last_list.append(pre_wk_last)
    else:
        print ('-----------非周五，不计算周度数据-----------')
        prewk_last_list = [''] * 12
    return prewk_last_list


def getSoup(url, kv):
    try:
        requests.adapters.DEFAULT_RETRIES = 5 # 增加连接重试次数
        s = requests.session()
        s.keep_alive = False # http connection关闭keep-alive
        # r = requests.get(url, headers = kv)
        r = requests.get(url, headers = kv, verify=False)
        r.raise_for_status()
        r.encoding = 'utf-8'
        # print("status_code: " + str(r.status_code))
        # print("encoding: " + r.encoding)
    except:
        print("\n\n爬取失败\n\n")
    demo = r.text
    soup = BeautifulSoup(demo, "html.parser")
    print ('\n\n防止爬取过快，休息1秒钟\n\n')
    time.sleep(1)
    return soup


def getIndice(soup, rtday, index, ref, mark, prewk_last_list):
    # url = 'https://cn.investing.com/indices/major-indices'
    print("==============\n\n读取美欧股指\n\n==============")
    phrase = ''
    phrase1 = ''
    phrase2 = ''
    close_date = ('0'+str(rtday.day))[-2:]+'/'+('0'+str(rtday.month))[-2:]
    for n in range(8):
        i = index[n]
        j = ref[n]
        m = mark[n]
        prewk_last = prewk_last_list[n]
        prewk_last = prewk_last.replace(',', '')
        last = soup.select("."+j+"-last")[0].text
        pcp = soup.select("."+j+"-pcp")[0].text
        trad_time = soup.select("."+j+"-time")[0].text
        pcp = pcp.strip('+')
        last = last.replace(',', '')
        if prewk_last == '':
            prewk_txt = ''      
        else:
            prewk_pcp = str(round((float(last)/float(prewk_last)-1)*100, 2))
            if '-' in prewk_pcp:
                prewk_zd = u'，周跌'
                prewk_pcp = prewk_pcp.strip('-')
            else:
                prewk_zd = u'，周涨'
            prewk_txt = prewk_zd+prewk_pcp+u'%'
        if '-' in pcp:
            zd = u'收跌'
            pcp = pcp.strip('-')
        else:
            zd = u'收涨'
        if n <= 2:
            if trad_time.find(':') != -1: 
                print (i + u'收盘时间：'+trad_time)
                ph = i+zd+pcp+u"，"+u"报"+last+u'点'+prewk_txt+m
                print (ph+'\n---------')
                phrase1 = phrase1 + ph
                if n == 2:
                    sz = phrase1.count(u'收涨')
                    sd = phrase1.count(u'收跌')
                    if sz == 0:
                        ph_t = u'美国三大股指全线下跌。'
                    elif sd == 0:
                        ph_t = u'美国三大股指全线上涨。'
                    else:
                        ph_t = u'美国三大股指涨跌互现。'
                    phrase1 = ph_t + phrase1
                else:                        
                    pass
            else:
                phrase1 = u'美国股市今日休市'
        else:
            if trad_time.find(':') != -1 or trad_time == close_date:
                print (i + u'收盘时间：'+trad_time)
                ph = i+zd+pcp+u"，"+u"报"+last+u'点'+prewk_txt+m
                print (ph+'\n---------')
                phrase2 = phrase2 + ph                   
            else:
                if n <= 5:
                    k = i[:2]
                else:
                    k = i[:3]
                ph = u'%s今日休市'%k+prewk_txt+m
                print (ph+'\n---------')
                phrase2 = phrase2 + ph
            if n == 7:
                sz = phrase2.count(u'收涨')
                sd = phrase2.count(u'收跌')
                if sz == 0:
                    ph_t = u'欧洲主要股指集体下跌。'
                elif sd == 0:
                    ph_t = u'欧洲主要股指集体上涨。'
                else:
                    ph_t = u'欧洲主要股指涨跌互现。'
                phrase2 = ph_t + phrase2
            else:
                pass
    phrase = phrase1 + phrase2
    print (phrase+'\n\n')
    return phrase


def getSector(soup, sector):
    # url = 'https://cn.tradingview.com/markets/indices/quotes-snp/'
    print("==============\n\n读取SP行业板块\n\n==============")
    sector_name = []
    sector_pcp = []
    t = soup.find_all('tr', {'class': 'tv-data-table__row tv-data-table__stroke tv-screener-table__result-row'})
    for i in t:
        j = i.attrs['data-symbol']
        k = sector[j]
        sector_name.append(k)
    for i in t:
        l = re.findall('[pn]\">.*%', str(i))
        ln = len(l)
        if ln > 1:
            print ('长度为%d')%ln
            l = l[0]
        l = l[0]
        lm = l.find('>')+1
        l = l[lm:]
        sector_pcp.append(l)
    result = pd.DataFrame({'name':sector_name, 'pcp':sector_pcp,})
    result.replace('%', '', inplace = True, regex=True)
    result['pcp'] = pd.to_numeric(result['pcp'])
    # result['pcp'] = result['pcp'].map(lambda x: x / 1)
    # result['pcp'] = round(result['pcp'], 2)
    result = result.sort_values(by=['pcp'], ascending = False)
    result = result.reset_index(drop = True)
    x = len(result[result['pcp']<0])
    phrase = ''
    phrase1 = ''
    phrase2 = ''
    if x == 0:
        ph_t = u'标普500行业板块全线上涨，其中'
        p = result.loc[:1]
        phrase = u'%s、%s板块领涨，分别上涨%s%%、%s%%。'%(p.iat[0,0], p.iat[1,0], p.iat[0,1], p.iat[1,1])
        phrase = ph_t + phrase
    elif x == 11:
        ph_t = u'标普500行业板块全线下跌，其中'
        p = result.loc[9:10]
        phrase = u'%s、%s板块领跌，分别下跌%s%%、%s%%。'%(p.iat[1,0], p.iat[0,0], abs(p.iat[1,1]), abs(p.iat[0,1]))
        phrase = ph_t + phrase       
    elif x < 3:
        ph_t = u'标普500行业板块多数上涨，其中'
        p1 = result.loc[:1]
        phrase1 = u'%s、%s板块领涨，分别上涨%s%%、%s%%；'%(p1.iat[0,0], p1.iat[1,0], p1.iat[0,1], p1.iat[1,1])
        p2 = result.loc[9:10]
        if x == 2:
            phrase2 = u'仅%s、%s板块分别下跌%s%%、%s%%。'%(p2.iat[1,0], p2.iat[0,0], abs(p2.iat[1,1]), abs(p2.iat[0,1]))
        else:
            phrase2 = u'仅%s板块下跌%s%%。'%(p2.iat[1,0], abs(p2.iat[1,1]))
        phrase = ph_t + phrase1 + phrase2
    elif x > 9:
        ph_t = u'标普500行业板块多数下跌，其中'
        p1 = result.loc[9:10]
        phrase1 = u'%s、%s板块领跌，分别下跌%s%%、%s%%；'%(p1.iat[1,0], p1.iat[0,0], abs(p1.iat[1,1]), abs(p1.iat[0,1]))
        p2 = result.loc[:1]
        if x == 2:
            phrase2 = u'仅%s、%s板块分别上涨%s%%、%s%%。'%(p2.iat[0,0], p2.iat[1,0], p2.iat[0,1], p2.iat[1,1])
        else:
            phrase2 = u'仅%s板块下跌%s%%。'%(p2.iat[1,0], p2.iat[1,1])
        phrase = ph_t + phrase1 + phrase2
    else:
        ph_t = u'标普500行业板块涨跌互现，其中'
        p1 = result.loc[:1]
        phrase1 = u'%s、%s板块领涨，分别上涨%s%%、%s%%；'%(p1.iat[0,0], p1.iat[1,0], p1.iat[0,1], p1.iat[1,1])
        p2 = result.loc[9:10]
        phrase2 = u'%s、%s板块领跌，分别下跌%s%%、%s%%。'%(p2.iat[1,0], p2.iat[0,0], abs(p2.iat[1,1]), abs(p2.iat[0,1]))
        phrase = ph_t + phrase1 + phrase2
    print (phrase+'\n\n')
    return phrase


def getVix(soup, index, ref, prewk_last_list):
    # url = 'https://cn.investing.com/indices/major-indices'
    print("==============\n\n读取VIX指数\n\n==============")
    i = index[8]
    j = ref[8]
    prewk_last = prewk_last_list[8]
    last = soup.select('.'+j+'-last')[0].text
    pcp = soup.select('.'+j+'-pcp')[0].text
    pcp = pcp.strip('+')
    last = last.replace(',', '')
    trad_time = soup.select('.'+j+'-time')[0].text
    print (u'收盘时间：'+trad_time)
    if trad_time.find('/') == -1 and trad_time.find(':') != -1:
        print (u'收盘时间：'+trad_time) 
        if prewk_last == '':
            prewk_txt = ''
        else:
            prewk_pcp = str(round((float(last)/float(prewk_last)-1)*100, 2))
            if '-' in prewk_pcp:
                prewk_zd = u'，周跌'
                prewk_pcp = prewk_pcp.strip('-')
            else:
                prewk_zd = u'，周涨'
            prewk_txt = prewk_zd+prewk_pcp+'%'
        if '-' in pcp:
            zd = u'收跌'
            pcp = pcp.strip('-')
            ph_t = u'震荡下行，'
        else:
            zd = u'收涨'
            ph_t = u'震荡上行，'
        phrase = i+ph_t+zd+pcp+u"，"+u"报"+last+prewk_txt+u'。'
    else:
        phrase = '%s今日休市。'%i
    print (phrase+'\n\n')
    return phrase


def getTB(soup_1, soup_2, TBlist, prewk_last_list):
    # url = 'https://cn.investing.com/rates-bonds/world-government-bonds'
    print("==============\n\n读取美国国债收益率\n\n==============")
    report = pd.DataFrame(columns=['name', 'last', 'pc'])
    t = soup_1.find_all('table', {'id': 'rates_bonds_table_1'})
    prewk_last = prewk_last_list[9]
    cnt = 0
    while cnt < 10:
        tb = TBlist[cnt]
        for i in t:                     
            j = i.find_all('tr', {'id': 'pair_'+tb})
            for k in j:
                tb_name = k.a.attrs['title']
                tb_last = k.find('td', {'class': 'pid-'+tb+'-last'}).text
                tb_pc = k.find('td', {'class': 'pid-'+tb+'-pc'}).text
            report.loc[cnt] = [tb_name, tb_last, tb_pc]
        cnt += 1
    for i in range(len(report['pc'])):
        report['pc'][i] = report['pc'][i].strip('+')        
    report['pc'] = pd.to_numeric(report['pc'])
        # result['pcp'] = result['pcp'].map(lambda x: x / 1)
        # result['pcp'] = round(result['pcp'], 2)
    x = len(report[report['pc']<0])
    if x == 0:
        ph_t = u'美债收益率全线上涨，其中10年期美债收益率'
    elif x == 10:
        ph_t = u'美债收益率全线下跌，其中10年期美债收益率'        
    elif x < 3:
        ph_t = u'美债收益率多数上涨，其中10年期美债收益率'
    elif x > 7:
        ph_t = u'美债收益率多数下跌，其中10年期美债收益率'         
    else:
        ph_t = u'美债收益率涨跌互现，其中10年期美债收益率'        
    # 获取10年期国债上一日实际收盘价
    t2 = soup_2.find_all('table', {'class': 'genTbl closedTbl historicalTbl'})
    for i in t2:
        j = i.text
    k = j.split()
    pre_last = k[13]
    last = report['last'][8]
    # pc = round(report['pc'][8]*100, 2)
    pc = round((float(last) - float(pre_last))*100, 2)
    if prewk_last == '':
        prewk_txt = ''
    else:
        prewk_pc = str(round((float(last)-float(prewk_last))*100, 2))
        if '-' in prewk_pc:
            prewk_zd = u'，周跌'
            prewk_pc = prewk_pc.strip('-')        
        else:
            prewk_zd = u'，周涨'
        prewk_txt = prewk_zd+prewk_pc+u'个基点'
    if pc > 0:
        zd = u'收涨'
        phrase = ph_t+zd+str(pc)+u'个基点，报%s%%'%last+prewk_txt+u'。'
    elif pc < 0:
        zd = u'收跌'
        pc = abs(pc)
        phrase = ph_t+zd+str(pc)+u'个基点，报%s%%'%last+prewk_txt+u'。'
    else:
        zd = u'收平'
        phrase = ph_t+zd+u'，报%s%%'%last+prewk_txt+'。'
    print (phrase+'\n\n')
    return phrase
            

def getUSD(soup, prewk_last_list):
    print("==============\n\n读取美元指数\n\n==============")
    # url = 'https://cn.investing.com/currencies/us-dollar-index'  
    prewk_last = prewk_last_list[10]
    last = soup.select('.pid-8827-last')[0].text
    pcp = soup.select('.pid-8827-pcp')[0].text
    pcp = pcp.strip('+')
    trad_time = soup.select('.pid-8827-time')[0].text
    print (u'收盘时间：'+trad_time)
    if trad_time.find('/') == -1 and trad_time.find(':') != -1:
        print ('判断大盘情况：正常运行')  
        if prewk_last == '':
            prewk_txt = ''
        else:
            prewk_pcp = str(round((float(last)/float(prewk_last)-1)*100, 2))
            if '-' in prewk_pcp:
                prewk_zd = u'，周跌'
                prewk_pcp = prewk_pcp.strip('-')
            else:
                prewk_zd = u'，周涨'
            prewk_txt = prewk_zd+prewk_pcp+'%'
        if '-' in pcp:
            zd = u'收跌'
            pcp = pcp.strip('-')
        else:
            zd = u'收涨'
        phrase = u'美元指数'+zd+pcp+u'，报'+last+prewk_txt+u'。'
    else:
        phrase = u'美元指数期货今日休市。'
    print (phrase+'\n\n')
    return phrase


def getA50(soup, prewk_last_list):
    print("==============\n\n读取富时A50期货\n\n==============")
    # url = 'https://cn.investing.com/indices/china-a50'
    last = soup.select('.pid-44486-last')[0].text
    pcp = soup.select('.pid-44486-pcp')[0].text
    last = last.replace(',', '')
    pcp = pcp.strip('+')
    prewk_last = prewk_last_list[11]
    prewk_last = prewk_last.replace(',', '')
    trad_time = soup.select('.pid-44486-time')[0].text
    print (u'收盘时间：'+trad_time)
    if trad_time.find('/') == -1:
        print ('判断大盘情况：正常运行')  
        if prewk_last == '':
            prewk_txt = ''
        else:
            prewk_pcp = str(round((float(last)/float(prewk_last)-1)*100, 2))
            if '-' in prewk_pcp:
                prewk_zd = u'，周跌'
                prewk_pcp = prewk_pcp.strip('-')
            else:
                prewk_zd = u'，周涨'
            prewk_txt = prewk_zd+prewk_pcp+'%'
        if '-' in pcp:
            zd = u'收跌'
            pcp = pcp.strip('-')
        else:
            zd = u'收涨'
        phrase = u'富时中国A50期指'+zd+pcp+u'，报'+last+u'点'+prewk_txt+u'。'
    else:
        print ('今日闭盘')
        phrase = u'富时中国A50期指今日休市。'
    print (phrase+'\n========')
    return phrase


if __name__ == '__main__':
    # 导入准备数据、昨日日期
    url_list, url_list_his, kv, index, ref, mark, sector, TBlist = getDict()
    rtday= dt.datetime.today()-dt.timedelta(days=1) 
    report_date = str(rtday.month)+'月'+str(rtday.day)+'日'
    report_title = report_date+'美欧市场情况简析'
    prewk_last_list = weekEnd(rtday, url_list_his, getSoup, kv)
    # 获取各项行情，汇总至txt
    print ('----------开始跑数，日报日期为%s--------------'%report_date)
    soup_index = getSoup(url_list[0], kv)
    txt_index = getIndice(soup_index, rtday, index, ref, mark, prewk_last_list)
    txt1_index_1 = txt_index[:txt_index.find(u'欧洲')]
    txt1_index_2 = txt_index[txt_index.find(u'欧洲'):]
    txt_vix = getVix(soup_index, index, ref, prewk_last_list)
    
    soup_sector = getSoup(url_list[1], kv)
    txt_sector = getSector(soup_sector, sector)
    
    soup_TB_1 = getSoup(url_list[2], kv)
    soup_TB_2 = getSoup(url_list_his[9], kv)
    txt_TB = getTB(soup_TB_1, soup_TB_2, TBlist, prewk_last_list)
    
    soup_USD = getSoup(url_list[3], kv)
    txt_USD = getUSD(soup_USD, prewk_last_list)
    
    soup_A50 = getSoup(url_list[4], kv)
    txt_A50 = getA50(soup_A50, prewk_last_list)
    
    txt_fin  = txt1_index_1 + txt_sector + txt1_index_2 + txt_vix + txt_TB + txt_USD + txt_A50
    print (txt_fin)
    print ('----------跑数完毕，在txt_fin中查看--------------')
    
    '''
    # 导出docx文件
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    doc = Document()
    h = doc.add_heading(report_title, 0)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # doc.add_heading('一、行情信息', level=1)
    p = doc.add_paragraph()
    run=p.add_run(txt_fin)
    run.font.name = u'宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.color.rgb = RGBColor(0,0,0)#设置颜色为黑色
    doc.save('%s.docx'%report_title)#保存到工程目录下
    print ('----------文件导出完毕，在《%s.docx》中查看--------------'%report_title)
    '''
