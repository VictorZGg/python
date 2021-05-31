# -*- coding: utf-8 -*-
"""
Created on Tue Feb 23 16:26:39 2021

@author: zengg
"""

import pandas as pd
import requests
from bs4 import BeautifulSoup
import re
import time
import datetime as dt
import urllib3
# import xlsxwriter


def getSoup(url, kv):
    try:
        requests.adapters.DEFAULT_RETRIES = 5 # 增加连接重试次数
        s = requests.session()
        s.keep_alive = False # http connection关闭keep-alive
        # r = requests.get(url, headers = kv)
        r = requests.get(url, headers = kv, verify=False)
        r.raise_for_status()
        r.encoding = 'utf-8'
        # print("status_code: " + str(r.status_code))
        # print("encoding: " + r.encoding)
    except:
        print("\n\n爬取失败\n\n")
    demo = r.text
    soup = BeautifulSoup(demo, "html.parser")
    # print ('\n\n防止爬取过快，休息1秒钟\n\n')
    time.sleep(0.1)
    return soup


def getContent(url_1, kv, end_year, page_no):
    report = pd.DataFrame(columns=['发布机构','发文日期','名称','文号','内容'])
    year = 2021
    for n in range(page_no):
        if year < end_year:
            break
            print('当前年数小于'+str(end_year)+'年，终止')
        else:
            print ('==========='+'第'+str(n+1)+'页===========')
            m = ['_'+str(n),''][n==0] # 判断n是否为0，决定m的取值
            url = url_1+m+'.htm'
            soup = getSoup(url, kv)
            t = soup.find_all('div', {'class': 'row'})
            for i in t:
                # 进入文书页面爬取
                ref = 'http://www.csrc.gov.cn/pub/zjhpublic/' + i.find('a')['href'][6:]
                soup2 = getSoup(ref, kv)
                # 查询标题信息
                t2 = soup2.find_all('div', {'class': 'headInfo'})
                j2 = t2[0].find_all(['span'])
                text = []
                for k in range(1, (len(j2)-1)):
                    x = (j2[k].text)
                    text.append(x) # 循环注入标题信息
                # 查询文书内容
                content = soup2.find('div', {'id': 'ContentRegion'}).text
                # 合并标题与内容
                print (text)
                print ('--------------------------')
                text.append(content)
                # 合并内容插入表格中
                report.loc[len(report)] = text
    return report


if __name__ == '__main__':
    url_1 = 'http://www.csrc.gov.cn/pub/zjhpublic/3300/3313/index_7401'
    kv = {'user-agent': 'Mozilla/5.0',
          'Connection':'close'}
    end_year = 2021 # 设置到哪一年为止
    page_no = 10 # 设置查多少页
    
    ###################### 获取文件列表 ######################
    x = getContent(url_1, kv, end_year, page_no)
    
    ###################### 提取当事人、违法事实、处罚信息 ######################
    keyword_y_list = ['男','女','住址','住所','有限公司','事务所']
    keyword_n_list = ['规定','调查','审理','事实','申辩']
    table = pd.DataFrame(columns=['发布机构','发文日期','名称','文号','当事人','违法事实','处罚信息'])
    # i = 0
    for i in range(len(x)):
        x1 = x['内容'][i]
        patter1 = re.compile('违法事实.*?\n')
        try:
            s = patter1.search(x1).group()
        except:
            pass
        split_1 = x1.find(s)
        split_2 = x1.find('决定：')
        
        print ('==================当事人==================')
        x2 = x1[:split_1]
        x2 = x2.split('\n')
        person_list = []
        for j1 in x2:
            try:
                judge1 = any(w in j1 and w for w in keyword_y_list)
                judge2 = any(w in j1 and w for w in keyword_n_list)
                if judge1 and not judge2:
                    print (j1)
                    person_list.append(j1)
            except:
                person_list.append('')
       
        print ('==================违法事实==================')
        x3 = x1[split_1:split_2]
        fact_list = []
        if any(k in x3 and k for k in ['一、','二、']):
            x3 = x3.split('\n')
            # j2 = x3[0]
            for j2 in x3:
                try:
                    if (j2[0] in ['一','二','三','四','五','（']) or (j2[1] == '.'):
                        print (j2)
                        fact_list.append(j2)
                except:
                    fact_list.append('')
        else:
            fact_list.append(x3)
        
        print ('==================处罚信息==================')
        x4 = x1[split_2:]
        penalty_list = []
        if any(k in x4 and k for k in ['一、','二、']):
            x4 = x4.split('\n')
            # j2 = x3[0]
            for j3 in x4:
                try:
                    if (j3[0] in ['一','二','三','四','五','（']) or (j3[1] == '.'):
                        print (j3)
                        penalty_list.append(j3)
                except:
                    penalty_list.append('')
        else:
            penalty_list.append(x4)
       
        person = '\n'.join(person_list)
        fact = '\n'.join(fact_list)
        penalty = '\n'.join(penalty_list)
        
        table_i = [x.iloc[i,0],x.iloc[i,1],x.iloc[i,2],x.iloc[i,3],person,fact,penalty]
        table.loc[len(table)] = table_i    
    
    ###################### 全部导出至excel文件 ######################
    path = 'D:\\tools\\python\\python\\zjh\\penalty_list.xlsx'
    table.to_excel(path,index=False,header=True)
    
    import os
    os.system(path)        
    
    
    ###################### 按照关键字导出至word文件 ######################
    from docx import Document
    from docx.oxml.ns import qn
    # from docx.enum.text import WD_ALIGN_PARAGRAPH
    # from docx.shared import RGBColor
    key_word = ['虚增', '虚减']
    
    for i in key_word:       
        index = x['内容'].str.contains(i)
        key_content = x[index]
        key_content = key_content.reset_index(drop=True)
        
        doc = Document()
        for j in range(len(key_content)):
            run = doc.add_heading('',level=1).add_run(key_content.loc[j,'名称'])
            # run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.name=u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
            doc.styles['Normal'].font.name = u'仿宋'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            p = doc.add_paragraph('发文日期：'+key_content.loc[j,'发文日期'])
            # p.add_run('文号：'+key_content.loc[1,'文号'])
            p.add_run(key_content.loc[j,'内容'])
        doc.save('%s.docx'%i)#保存到工程目录下
        print ('----------文件导出完毕，在《%s.docx》中查看--------------'%i)
    
    
    ###################### 全部导出至excel文件 ######################
    # x.to_excel('./'+'虚增'+'test.xlsx',index=False,header=True)
    
        
    
        
