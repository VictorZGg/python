# -*- coding: utf-8 -*-
"""
Created on Tue Feb 23 16:26:39 2021

@author: zengg
"""

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import os
print (os.getcwd())
os.chdir('D:\\tools\\python\\python\\zjh')


origin_table = pd.read_excel(io = 'enforcement_table.xlsx')
# table_len = len(origin_table)
# i = key_word[0]
# j = 0
doc = Document()

key_word = ['操纵', '内幕信息', '未公开信息']
for i in key_word:
    print ('--------- 关键词：' + i + '-----------')
    index = origin_table['内容'].str.contains(i)
    key_content = origin_table[index]
    key_content = key_content.reset_index(drop=True)
    table_len = len(key_content)
    for j in range(table_len):
        print ('=========== 进度：'+str(round((j+1)/table_len*100,2))+'% ===========')
        table_i = key_content.loc[j]
        run = doc.add_heading('',level=1).add_run(table_i['发文日期']+'——'+table_i['名称'])
        # run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.name=u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
        doc.styles['Normal'].font.name = u'仿宋'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        p = doc.add_paragraph('发文日期：'+table_i['发文日期'])
        # p.add_run('文号：'+report.loc[1,'文号'])
        p.add_run(table_i['内容'])
    doc.save('处罚决定书：'+i+'.docx')#保存到工程目录下
print ('----------文件导出完毕--------------')

    
        
    
        
